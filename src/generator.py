"""
The password generator app.

A small application that allows you to generate a password of any complexity using a user-friendly
graphical interface with accessible settings and save it in a convenient 'txt' file so you don't forget
your social media passwords.

This module contains the core logic for generating password, evaluating their strength,
and it's used by both the CLI and the GUI versions of the application.
"""
import json
from secrets import choice
from docx import Document
import string
import time
import os
import csv


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PASSWORDS_FILE = os.path.join(BASE_DIR, "settings.json")
SETTINGS_FILE = os.path.join(BASE_DIR, "settings.json")


# ===============================================================================================
# 1. AUXILIARY FUNCTIONS
# ===============================================================================================


def greed():
    """
    Display welcome message.

    Shows a formatted welcome screen with emojis and explains the goal
    of the game. Includes short pauses for better user experience.
    """
    print("\n" + "=" * 50)
    print("🔑 Welcome to the 'Personal Password Generator'!")
    print("=" * 50)

    time.sleep(0.5)
    print("🛡️ You can create a password of the developer's specified "
          "length (from 16 to 64) and flexibly customize it to suit "
          "your needs!")


def get_number(prompt):
    """
       Get a valid integer from the user.

       Continuously prompts the user until a valid integer is entered.
       Handles ValueError exceptions and displays an error message.
       """
    while True:
        try:
            user_length = int(input(prompt))
            if user_length < 16:
                print("⚠️ Warning! The your password is too short!")
                print("🛡️ The password must be at least 16 characters long!")
                continue
            if user_length > 64:
                print("⚠️ Warning! The password is too long!")
                print("🛡️ The password must be at least 64 characters long!")
                continue
            return user_length
        except ValueError:
            print("❌ Error! Please enter a number!")


# ==============================================================================================
# 2. APP LOGIC
# ==============================================================================================


def get_character_options():
    """
    Get password complexity settings from the user
    based on the presence of letters, numbers, and symbols.

    Returns data for creating a password.
    """
    user_letters = input("Include letters? (y/n): ").lower() == 'y'
    user_digits = input("Include digits? (y/n): ").lower() == 'y'
    user_symbols = input("Include symbols? (y/n): ").lower() == 'y'
    return user_letters, user_digits, user_symbols


def build_character_pool(user_letters, user_digits, user_symbols):
    """
    Build a character pool based on user preferences.

    Returns a string containing allowed characters for password generation.
    """
    characters = ""
    if user_letters:
        characters += string.ascii_letters
    if user_digits:
        characters += string.digits
    if user_symbols:
        characters += string.punctuation
    return characters


def generate_password(length, character_pool):
    """
    Generate a password based on the entered data.

    Generation will not occur if the user has not previously
    selected password settings (no letters, numbers, or symbols).
    """
    if not character_pool:
        print("❌ Error! No characters selected!")
        return None
    return "".join(choice(character_pool) for _ in range(length))


def check_strength(password):
    """
    Password strength check based on input data.
    Total password length, including presence of letters,
    numbers, and symbols.

    """
    score = 0

    # Check length
    if 16 <= len(password) <= 29:
        score += 1
    if len(password) >= 30:
        score += 2

    # Check all symbols
    if any(c.islower() for c in password):
        score += 1
    if any(c.isupper() for c in password):
        score += 1
    if any(c.isdigit() for c in password):
        score += 1
    if any(c in "!@#$%^&*()" for c in password):
        score += 1

    # Return result
    if score <= 2:
        return "Not Safe"
    elif score <= 4:
        return "Moderate"
    else:
        return "Very Strong"


def generate_again():
    """
    Asks the user if he wants to repeat the generation of a new password.
    """
    user_answer = input("Do you want to generate another password? (y/n): ").lower()
    positive = ['y', 'yes', 'yeah', 'da', 'd', 'a']
    negative = ['n', 'no', 'not', 'net', 'nope']

    if user_answer in positive:
        return True
    elif user_answer in negative:
        return False
    else:
        print("⚠️ Please answer 'y' or 'n'!")
        return generate_again()


def save_passwords_txt(file_path, passwords):
    with open(file_path, "w", encoding="utf-8") as f:
        for p in passwords:
            f.write(f"Service: {p['service']} | Login/email: {p['login']} | Password: {p['password']}\n")


def load_passwords_txt(file_path=PASSWORDS_FILE):
    if not os.path.exists(file_path):
        return []

    passwords = []
    with open(file_path, 'r', encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue

            parts = line.split(" | ")
            if len(parts) != 3:
                continue

            service = parts[0].replace("Service: ", "")
            login = parts[1].replace("Login/email: ", "")
            password = parts[2].replace("Password: ", "")

            passwords.append({
                "service": service,
                "login": login,
                "password": password
            })

    return passwords


def save_passwords_csv(file_path, passwords):
    with open(file_path, 'w', encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["service", "login", "password"])
        writer.writeheader()
        writer.writerows(passwords)


def load_passwords_csv(file_path):
    passwords = []
    with open(file_path, 'r', encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            passwords.append({
                "service": row.get("service", ""),
                "login": row.get("login", ""),
                "password": row.get("password", "")
            })
    return passwords


def save_passwords_json(file_path, passwords):
    with open(file_path, 'w', encoding="utf-8") as f:
        json.dump(passwords, f, ensure_ascii=False, indent=4)


def load_passwords_json(file_path):
    if not os.path.exists(file_path) or os.path.getsize(file_path) == 0:
        return []
    with open(file_path, 'r', encoding="utf-8") as f:
        return json.load(f)


def save_passwords_docx(file_path, passwords):
    doc = Document()
    doc.add_heading("Saved Passwords", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Service"
    hdr[1].text = "Login / Email"
    hdr[2].text = "Password"

    for p in passwords:
        row = table.add_row().cells
        row[0].text = p["service"]
        row[1].text = p["login"]
        row[2].text = p["password"]

    doc.save(file_path)


def save_passwords(file_path, passwords):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".csv":
        return save_passwords_csv(file_path, passwords)
    elif ext == ".json":
        return save_passwords_json(file_path, passwords)
    else:
        return save_passwords_txt(file_path, passwords)


def load_passwords(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".csv":
        return load_passwords_csv(file_path)
    elif ext == ".json":
        return load_passwords_json(file_path)
    else:
        return load_passwords_txt(file_path)


def search_passwords(password, query):
    pass





# ============================================================================================
# 3. MAIN FUNC
# ============================================================================================


def main():
    """
    The main controller of the application.

    Greets the user and starts the application cycle. After each cycle,
    asks the user if they want to restart the generation.
    """
    while True:
        greed()
        user_length = get_number("Enter your password length: ")

        print("\nChoose character types: ")
        user_letters, user_digits, user_symbols = get_character_options()

        character_pool = build_character_pool(user_letters, user_digits, user_symbols)

        if not character_pool:
            print("❌ You must choose at least one character type!")
            continue

        user_password = generate_password(user_length, character_pool)

        print("\n" + "=" * 50)
        print(f"✅ Your generated password: {user_password}")
        print("=" * 50)

        print(check_strength(user_password))

        if not generate_again():
            print("\n" + "=" * 55)
            print("👋 Thank you for using my Password Generator! Goodbye!")
            print("=" * 55)
            break

        else:
            print("\nReloading..", end="", flush=True)
            for _ in range(3):
                time.sleep(1)
                print(".", end="", flush=True)


# If you want to run the code directly, uncomment the lines below.
# In this case, you will not have access to the graphical interface,
# its settings, or the file-saving feature.

# if __name__ == "__main__":
#     main()